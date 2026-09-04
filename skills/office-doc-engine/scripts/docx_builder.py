"""MemoBuilder: legal memo/docx builder with real Word footnotes and AmLaw-grade formatting.

Permanent formatting spec (ben-writing-style-2/references/docx-formatting.md):
  1. 10.5pt body and headings (except title)
  2. 13.5pt title
  3. Times New Roman everywhere
  4. All text black (Automatic)
  5. Footnotes 8pt
  6. Real Word footnotes (native footnotes part)
  7. Zero space before/after footnotes
  8. Body text fully justified
  9. Title in header, right-justified, 8pt italic, from page 2 onward
 10. Footer page numbers centered: "X of Y" (PAGE of NUMPAGES), 8pt
 11. Summary metadata: Author = "Benjamin Snipes", nothing else
 12. Contracts: signature pages, exhibits, schedules, appendices each start on own page
 13. AmLaw 100 professionalism
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import CONTENT_TYPE as CT, RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.opc.part import Part
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor

FONT = "Times New Roman"
BODY_PT = 10.5
TITLE_PT = 13.5
FOOTNOTE_PT = 8
HEADER_PT = 8
AUTHOR = "Benjamin Snipes"


def _el(tag, **attrs):
    e = OxmlElement(tag)
    for k, v in attrs.items():
        e.set(qn(k), str(v))
    return e


class _FootnotesPart(Part):
    def __init__(self, partname, content_type, element, package):
        self._element = element
        super().__init__(partname, content_type, package=package)

    @property
    def blob(self):
        from lxml import etree
        return etree.tostring(self._element, xml_declaration=True, encoding="UTF-8", standalone=True)


class MemoBuilder:
    def __init__(self, title=None, author=None, date=None, header_text=None):
        self.document = Document()
        self._footnote_id = 1
        self._footnotes = []  # list of (id, text)
        # header shows a short summary of the document, not necessarily the literal title
        self._title = header_text or title
        self._setup_base_style()
        self._setup_page()
        self._setup_metadata()
        if title:
            self._title_block(title, author, date)

    # ---------- setup ----------

    def _setup_metadata(self):
        """Set Summary metadata: Author = Benjamin Snipes, nothing else."""
        cp = self.document.core_properties
        cp.author = AUTHOR
        cp.last_modified_by = ""
        cp.comments = ""
        cp.subject = ""
        cp.keywords = ""
        cp.category = ""
        cp.title = ""

    def _setup_base_style(self):
        st = self.document.styles["Normal"]
        st.font.name = FONT
        st.font.size = Pt(BODY_PT)
        st.font.color.rgb = RGBColor(0, 0, 0)
        st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        pf = st.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(6)
        pf.line_spacing = 1.15
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for name in ("Heading 1", "Heading 2", "Heading 3", "Heading 4"):
            h = self.document.styles[name]
            h.font.name = FONT
            h.font.size = Pt(BODY_PT)
            h.font.bold = True
            h.font.color.rgb = RGBColor(0, 0, 0)
            h.paragraph_format.space_before = Pt(10)
            h.paragraph_format.space_after = Pt(4)
            h.paragraph_format.keep_with_next = True

    def _setup_page(self):
        sec = self.document.sections[0]
        sec.top_margin = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin = Inches(1.25)
        sec.right_margin = Inches(1.25)
        sec.different_first_page_header_footer = True
        self._build_header(sec)
        self._build_footer(sec)

    def _header_par(self, container):
        p = container.paragraphs[0] if container.paragraphs else container.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(self._title or "")
        run.font.name = FONT
        run.font.size = Pt(HEADER_PT)
        run.font.italic = True
        return p

    def _build_header(self, sec):
        # first-page header: empty
        fh = sec.first_page_header
        fh.is_linked_to_previous = False
        if fh.paragraphs:
            fh.paragraphs[0].text = ""
        # default header (page 2+): title, right, 9pt italic
        h = sec.header
        h.is_linked_to_previous = False
        self._header_par(h)

    def _footer_page_field(self, container):
        p = container.paragraphs[0] if container.paragraphs else container.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

        def add_field(instr):
            r = p.add_run()
            r.font.name = FONT
            r.font.size = Pt(HEADER_PT)
            fld1 = _el("w:fldChar", **{"w:fldCharType": "begin"})
            instrText = OxmlElement("w:instrText")
            instrText.set(qn("xml:space"), "preserve")
            instrText.text = instr
            fld2 = _el("w:fldChar", **{"w:fldCharType": "end"})
            r._r.append(fld1)
            r._r.append(instrText)
            r._r.append(fld2)

        def add_text(txt):
            r = p.add_run(txt)
            r.font.name = FONT
            r.font.size = Pt(HEADER_PT)

        add_field(" PAGE ")
        add_text(" of ")
        add_field(" NUMPAGES ")

    def _build_footer(self, sec):
        ff = sec.first_page_footer
        ff.is_linked_to_previous = False
        self._footer_page_field(ff)
        f = sec.footer
        f.is_linked_to_previous = False
        self._footer_page_field(f)

    def _title_block(self, title, author, date):
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(title)
        r.font.name = FONT
        r.font.size = Pt(TITLE_PT)
        r.font.bold = True
        if author or date:
            p2 = self.document.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p2.paragraph_format.space_after = Pt(10)
            r2 = p2.add_run("  |  ".join(x for x in (author, date) if x))
            r2.font.name = FONT
            r2.font.size = Pt(BODY_PT)

    # ---------- footnotes ----------

    def _ensure_footnotes_part(self):
        if getattr(self, "_footnotes_part_ready", False):
            return
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls

        xml = (
            '<w:footnotes %s>'
            '<w:footnote w:type="separator" w:id="-1"><w:p><w:pPr><w:spacing w:before="0" w:after="0"/></w:pPr><w:r><w:separator/></w:r></w:p></w:footnote>'
            '<w:footnote w:type="continuationSeparator" w:id="0"><w:p><w:pPr><w:spacing w:before="0" w:after="0"/></w:pPr><w:r><w:continuationSeparator/></w:r></w:p></w:footnote>'
            "</w:footnotes>"
        ) % nsdecls("w")
        element = parse_xml(xml)
        partname = PackURI("/word/footnotes.xml")
        part = _FootnotesPart(partname, CT.WML_FOOTNOTES, element, self.document.part.package)
        self.document.part.relate_to(part, RT.FOOTNOTES)
        self._footnotes_part = part
        self._footnotes_element = element
        self._footnotes_part_ready = True

    def add_footnote(self, paragraph, text):
        self._ensure_footnotes_part()
        fid = self._footnote_id
        self._footnote_id += 1

        # reference mark in body
        r = paragraph.add_run()
        r.font.name = FONT
        rPr = r._r.get_or_add_rPr()
        va = _el("w:vertAlign", **{"w:val": "superscript"})
        rPr.append(va)
        ref = _el("w:footnoteReference", **{"w:id": str(fid)})
        r._r.append(ref)

        # footnote body: 8pt, zero spacing, justified off (left)
        fn = _el("w:footnote", **{"w:id": str(fid)})
        p = _el("w:p")
        pPr = _el("w:pPr")
        spacing = _el("w:spacing", **{"w:before": "0", "w:after": "0"})
        pPr.append(spacing)
        pStyle = _el("w:pStyle", **{"w:val": "FootnoteText"})
        pPr.append(pStyle)
        p.append(pPr)
        # footnote ref mark inside footnote
        r1 = _el("w:r")
        rPr1 = _el("w:rPr")
        rPr1.append(_el("w:rStyle", **{"w:val": "FootnoteReference"}))
        r1.append(rPr1)
        r1.append(_el("w:footnoteRef"))
        p.append(r1)
        # text run, 8pt TNR
        r2 = _el("w:r")
        rPr2 = _el("w:rPr")
        rf = _el("w:rFonts", **{"w:ascii": FONT, "w:hAnsi": FONT})
        sz = _el("w:sz", **{"w:val": str(int(FOOTNOTE_PT * 2))})
        szCs = _el("w:szCs", **{"w:val": str(int(FOOTNOTE_PT * 2))})
        rPr2.append(rf)
        rPr2.append(sz)
        rPr2.append(szCs)
        r2.append(rPr2)
        t = _el("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = " " + text
        r2.append(t)
        p.append(r2)
        fn.append(p)
        self._footnotes_element.append(fn)
        return fid

    # ---------- content ----------

    def heading(self, text, level=1):
        return self.document.add_heading(text, level=level)

    def paragraph(self, text, bold=False, italic=False, style=None):
        p = self.document.add_paragraph(style=style)
        r = p.add_run(text)
        r.font.name = FONT
        r.font.size = Pt(BODY_PT)
        r.bold = bold
        r.italic = italic
        return p

    def paragraph_with_footnote(self, text, footnote, **kwargs):
        p = self.paragraph(text, **kwargs)
        self.add_footnote(p, footnote)
        return p

    def footnote_here(self, paragraph, footnote):
        return self.add_footnote(paragraph, footnote)

    def bullet(self, text):
        p = self.document.add_paragraph(style="List Bullet")
        r = p.add_run(text)
        r.font.name = FONT
        r.font.size = Pt(BODY_PT)
        return p

    def numbered(self, text):
        p = self.document.add_paragraph(style="List Number")
        r = p.add_run(text)
        r.font.name = FONT
        r.font.size = Pt(BODY_PT)
        return p

    def block_quote(self, text):
        p = self.document.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.right_indent = Inches(0.5)
        r = p.add_run(text)
        r.font.name = FONT
        r.font.size = Pt(BODY_PT)
        return p

    def rule(self):
        p = self.document.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = _el("w:pBdr")
        bottom = _el("w:bottom", **{"w:val": "single", "w:sz": "6", "w:space": "1", "w:color": "000000"})
        pBdr.append(bottom)
        pPr.append(pBdr)
        return p

    def table(self, rows, cols, header=None):
        t = self.document.add_table(rows=rows, cols=cols)
        t.style = "Table Grid"
        if header:
            for j, h in enumerate(header):
                cell = t.cell(0, j)
                cell.text = ""
                p = cell.paragraphs[0]
                r = p.add_run(h)
                r.bold = True
                r.font.name = FONT
                r.font.size = Pt(BODY_PT)
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.name = FONT
                        r.font.size = Pt(BODY_PT)
        return t

    def page_numbers_footer(self):
        # already built in _setup_page; kept for backward compatibility
        pass

    def page_break(self):
        """Insert a page break. Use before signature pages, exhibits, schedules, appendices."""
        from docx.enum.text import WD_BREAK
        p = self.document.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run()
        run.add_break(WD_BREAK.PAGE)
        return p

    def set_margins(self, inches=1.0):
        for s in self.document.sections:
            s.top_margin = Inches(inches)
            s.bottom_margin = Inches(inches)
            s.left_margin = Inches(inches)
            s.right_margin = Inches(inches)

    def save(self, path):
        self.document.save(path)
        return path
