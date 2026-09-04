"""
Redline Builder - Native Word Tracked Changes
==============================================
Builds a .docx with native Word tracked changes (w:ins/w:del) and comments
from a structured redline analysis (SECTION/ORIGINAL/PROPOSED/REASON entries).

Usage:
    python3 build_redline.py --source /path/to/original.pdf --entries /path/to/entries.json --output /path/to/output.docx [--author "Name"] [--cover-note /path/to/cover.txt]

The builder:
1. Extracts full text from the source document (PDF or DOCX)
2. Matches each entry's ORIGINAL text to its position in the document
3. Builds a Word document with:
   - Cover note (first page)
   - Full agreement text with tracked changes at matched positions
   - Native Word comments with rationale
4. Verifies the output before saving

This produces a .docx that opens in Microsoft Word with tracked changes
displayed exactly as if a human lawyer made them in Word.
"""

import sys
import os
import re
import json
import zipfile
from datetime import datetime
from pathlib import Path

# Add skill scripts to path
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))


def extract_text(source_path: str) -> str:
    """Extract full text from PDF or DOCX."""
    suffix = Path(source_path).suffix.lower()
    
    if suffix == '.pdf':
        import pdfplumber
        pages = []
        with pdfplumber.open(source_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text() or ""
                pages.append(text)
        full_text = "\n".join(pages)
        # Clean up page artifacts
        full_text = re.sub(r'\nConfidential\s*\n', '\n', full_text)
        full_text = re.sub(r'\n\s*\d+\s*\n', '\n', full_text)
        return full_text
    
    elif suffix in ('.docx', '.doc'):
        from docx import Document
        doc = Document(source_path)
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    
    elif suffix in ('.txt', '.md'):
        return Path(source_path).read_text()
    
    else:
        raise ValueError(f"Unsupported file type: {suffix}")


def normalize(text: str) -> str:
    """Normalize text for matching."""
    text = re.sub(r'\s+', ' ', text).strip()
    text = text.replace('\u2019', "'").replace('\u2018', "'")
    text = text.replace('\u201c', '"').replace('\u201d', '"')
    text = text.replace('\u2013', '-').replace('\u2014', '-')
    text = text.replace('\u2026', '...')
    return text


def find_matches(full_text: str, entries: list) -> list:
    """Find position of each entry's ORIGINAL text in the full document."""
    full_norm = normalize(full_text)
    matches = []
    
    for idx, entry in enumerate(entries):
        orig_norm = normalize(entry['original'])
        
        # Try exact match
        pos = full_norm.find(orig_norm)
        if pos < 0 and len(orig_norm) > 30:
            # Try first 30 chars
            pos = full_norm.find(orig_norm[:30])
        
        if pos >= 0:
            end_pos = pos + len(orig_norm)
            matches.append((pos, end_pos, idx))
    
    # Sort by position and remove overlaps
    matches.sort(key=lambda x: x[0])
    filtered = []
    last_end = 0
    for start, end, idx in matches:
        if start >= last_end:
            filtered.append((start, end, idx))
            last_end = end
    
    return filtered


def build_docx(full_text: str, entries: list, matches: list, 
               output_path: str, author: str = "Benjamin Snipes",
               cover_note: str = None):
    """Build the .docx with tracked changes and comments."""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    DATE = datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ")
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10.5)
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Cover note
    if cover_note:
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run("COVER NOTE")
        run.bold = True
        run.font.size = Pt(13.5)
        run.font.name = 'Times New Roman'
        
        doc.add_paragraph()
        
        for line in cover_note.split('\n'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run(line)
            run.font.size = Pt(10.5)
            run.font.name = 'Times New Roman'
        
        doc.add_page_break()
    
    # Helper functions for OOXML elements
    def make_run(text):
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '21')  # 10.5pt
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), '21')
        rPr.append(szCs)
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rPr.append(rFonts)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        r.append(t)
        return r
    
    def make_del(text, rev_id):
        del_elem = OxmlElement('w:del')
        del_elem.set(qn('w:id'), str(rev_id))
        del_elem.set(qn('w:author'), author)
        del_elem.set(qn('w:date'), DATE)
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        r.append(rPr)
        dt = OxmlElement('w:delText')
        dt.set(qn('xml:space'), 'preserve')
        dt.text = text
        r.append(dt)
        del_elem.append(r)
        return del_elem
    
    def make_ins(text, rev_id):
        ins_elem = OxmlElement('w:ins')
        ins_elem.set(qn('w:id'), str(rev_id))
        ins_elem.set(qn('w:author'), author)
        ins_elem.set(qn('w:date'), DATE)
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        r.append(rPr)
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        r.append(t)
        ins_elem.append(r)
        return ins_elem
    
    # Build the document body with tracked changes
    # Split into paragraphs
    full_norm = normalize(full_text)
    paragraphs = re.split(r'\n\s*\n', full_text)
    paragraphs = [p.strip() for p in paragraphs if p.strip()]
    
    rev_id = 100
    comment_id = 0
    comments_data = []
    
    # Track which matches have been used
    used_matches = set()
    
    for para_text in paragraphs:
        para_norm = normalize(para_text)
        if not para_norm:
            continue
        
        # Find this paragraph in the full normalized text
        para_pos = full_norm.find(para_norm)
        if para_pos < 0 and len(para_norm) > 50:
            para_pos = full_norm.find(para_norm[:50])
        
        if para_pos < 0:
            # Just add as regular paragraph
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run(para_text.replace('\n', ' '))
            run.font.size = Pt(10.5)
            run.font.name = 'Times New Roman'
            continue
        
        # Find matches within this paragraph
        para_end = para_pos + len(para_norm)
        para_matches = [(s, e, i) for s, e, i in matches 
                       if s >= para_pos and s < para_end and i not in used_matches]
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        if para_matches:
            pos_in_para = 0
            for start, end, idx in para_matches:
                entry = entries[idx]
                used_matches.add(idx)
                
                rel_start = start - para_pos
                rel_end = min(end - para_pos, len(para_norm))
                
                # Text before the match
                before = para_norm[pos_in_para:rel_start]
                if before.strip():
                    p._p.append(make_run(before))
                
                # Comment range start
                crs = OxmlElement('w:commentRangeStart')
                crs.set(qn('w:id'), str(comment_id))
                p._p.append(crs)
                
                # Deletion
                p._p.append(make_del(entry['original'], rev_id))
                rev_id += 1
                
                # Insertion
                p._p.append(make_ins(entry['proposed'], rev_id))
                rev_id += 1
                
                # Comment range end
                cre = OxmlElement('w:commentRangeEnd')
                cre.set(qn('w:id'), str(comment_id))
                p._p.append(cre)
                
                # Comment reference
                r = OxmlElement('w:r')
                rPr = OxmlElement('w:rPr')
                r.append(rPr)
                cr = OxmlElement('w:commentReference')
                cr.set(qn('w:id'), str(comment_id))
                r.append(cr)
                p._p.append(r)
                
                # Store comment
                comments_data.append((comment_id, entry.get('section', '?'), entry.get('reason', '')))
                comment_id += 1
                
                pos_in_para = rel_end
            
            # Text after last match
            after = para_norm[pos_in_para:]
            if after.strip():
                p._p.append(make_run(after))
        else:
            run = p.add_run(para_norm)
            run.font.size = Pt(10.5)
            run.font.name = 'Times New Roman'
    
    # Save base document
    base_path = output_path + ".base.docx"
    doc.save(base_path)
    
    # Build comments XML
    comments_xml = '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
    comments_xml += '<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">\n'
    
    for cid, section, reason in comments_data:
        reason_clean = reason.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('"', '&quot;')
        section_clean = str(section).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        comments_xml += f'''  <w:comment w:id="{cid}" w:author="{author}" w:date="{DATE}" w:initials="BS">
    <w:p>
      <w:pPr><w:pStyle w:val="CommentText"/></w:pPr>
      <w:r>
        <w:rPr><w:rStyle w:val="CommentReference"/></w:rPr>
        <w:annotationRef/>
      </w:r>
      <w:r>
        <w:t xml:space="preserve">Section {section_clean}: {reason_clean}</w:t>
      </w:r>
    </w:p>
  </w:comment>\n'''
    
    comments_xml += '</w:comments>'
    
    # Build final docx with comments injected
    with zipfile.ZipFile(base_path, 'r') as zin:
        with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == '[Content_Types].xml':
                    content = data.decode('utf-8')
                    if 'comments.xml' not in content:
                        content = content.replace('</Types>',
                            '<Override PartName="/word/comments.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"/>\n</Types>')
                    data = content.encode('utf-8')
                elif item.filename == 'word/_rels/document.xml.rels':
                    content = data.decode('utf-8')
                    if 'comments.xml' not in content:
                        content = content.replace('</Relationships>',
                            '<Relationship Id="rIdComments" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments" Target="comments.xml"/>\n</Relationships>')
                    data = content.encode('utf-8')
                zout.writestr(item, data)
            zout.writestr('word/comments.xml', comments_xml)
    
    # Clean up base file
    os.remove(base_path)
    
    return len(comments_data)


def main():
    import argparse
    
    parser = argparse.ArgumentParser(description="Build .docx redline with native tracked changes")
    parser.add_argument("--source", required=True, help="Path to original document (PDF/DOCX)")
    parser.add_argument("--entries", required=True, help="Path to redline entries JSON")
    parser.add_argument("--output", required=True, help="Output .docx path")
    parser.add_argument("--author", default="Benjamin Snipes", help="Tracked changes author")
    parser.add_argument("--cover-note", help="Path to cover note text file")
    
    args = parser.parse_args()
    
    # Load entries
    with open(args.entries) as f:
        entries = json.load(f)
    
    print(f"Source: {args.source}")
    print(f"Entries: {len(entries)}")
    print(f"Output: {args.output}")
    print(f"Author: {args.author}")
    
    # Extract text
    print("\nExtracting text...")
    full_text = extract_text(args.source)
    print(f"  Extracted: {len(full_text)} chars")
    
    # Find matches
    print("\nMatching entries to document...")
    matches = find_matches(full_text, entries)
    print(f"  Matched: {len(matches)}/{len(entries)}")
    
    if len(matches) < len(entries):
        unmatched = set(range(len(entries))) - set(idx for _, _, idx in matches)
        print(f"  Unmatched entries: {sorted(unmatched)}")
    
    # Load cover note
    cover_note = None
    if args.cover_note and os.path.exists(args.cover_note):
        cover_note = Path(args.cover_note).read_text()
    
    # Build document
    print("\nBuilding .docx with tracked changes...")
    comment_count = build_docx(full_text, entries, matches, args.output, 
                               author=args.author, cover_note=cover_note)
    
    print(f"  Comments created: {comment_count}")
    print(f"  File size: {os.path.getsize(args.output)} bytes")
    
    # Verify
    print("\nVerifying output...")
    from verify_redline import verify
    success = verify(args.output, args.entries)
    
    if success:
        print(f"\n✓ Redline built successfully: {args.output}")
    else:
        print(f"\n✗ Verification failed. Review errors above.")
        sys.exit(1)


if __name__ == "__main__":
    main()
