"""
Document RAG Pipeline for Contract Analysis
============================================
Layout-aware PDF extraction + vector store for targeted retrieval.

Usage:
    # Ingest a document
    python3 doc_rag.py ingest /path/to/document.pdf --name "BytePlus Agreement"
    
    # Query the store
    python3 doc_rag.py query "limitation of liability cap" --top_k 5
    
    # Get a specific section by number
    python3 doc_rag.py section "23.4"
    
    # List all ingested documents
    python3 doc_rag.py list
    
    # Delete a document
    python3 doc_rag.py delete "BytePlus Agreement"

The pipeline:
1. Extracts text with pdfplumber (preserves layout, detects tables)
2. Chunks by clause/section boundaries (legal-aware chunking)
3. Stores in chromadb with metadata (section number, page, heading)
4. Supports semantic query + exact section lookup
"""

import sys
import os
import re
import json
import hashlib
from pathlib import Path

STORE_PATH = "/home/user/.workspace/agent/doc_rag_store"


def extract_pdf(pdf_path: str) -> dict:
    """Extract text from PDF with layout awareness."""
    import pdfplumber
    
    pages = []
    with pdfplumber.open(pdf_path) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            tables = page.extract_tables()
            
            # Convert tables to text representation
            table_texts = []
            for table in tables:
                rows = []
                for row in table:
                    cells = [str(cell).strip() if cell else "" for cell in row]
                    rows.append(" | ".join(cells))
                table_texts.append("\n".join(rows))
            
            pages.append({
                "page_num": i + 1,
                "text": text,
                "tables": table_texts,
            })
    
    return {"pages": pages, "total_pages": len(pages)}


def extract_docx(docx_path: str) -> dict:
    """Extract text from DOCX."""
    from docx import Document
    
    doc = Document(docx_path)
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)
    
    # Also extract tables
    table_texts = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        table_texts.append("\n".join(rows))
    
    return {
        "pages": [{"page_num": 1, "text": "\n".join(paragraphs), "tables": table_texts}],
        "total_pages": 1,
    }


def chunk_by_clauses(text: str, doc_name: str) -> list:
    """
    Chunk legal document text by clause/section boundaries.
    Recognizes patterns like:
    - "1.1", "2.3.4", "12.2.1" (numbered clauses)
    - "SECTION 5", "ARTICLE III"
    - "Schedule 1", "Exhibit A", "Appendix B"
    - Headings (short lines followed by longer content)
    """
    chunks = []
    
    # Split into lines
    lines = text.split("\n")
    
    # Detect clause boundaries
    clause_pattern = re.compile(
        r'^(\d+(?:\.\d+)*(?:\([a-z0-9]+\))?)\s+'  # 1.1, 2.3(a), etc.
        r'|^(SECTION|ARTICLE|CLAUSE)\s+(\d+|[IVXLC]+)'  # SECTION 5, ARTICLE III
        r'|^(Schedule|Exhibit|Appendix|Annex|Supplement)\s+(\d+|[A-Z])',  # Schedule 1
        re.IGNORECASE
    )
    
    current_chunk = []
    current_section = "Preamble"
    current_start_line = 0
    
    for i, line in enumerate(lines):
        stripped = line.strip()
        if not stripped:
            continue
        
        match = clause_pattern.match(stripped)
        if match and len(current_chunk) > 0:
            # Save previous chunk
            chunk_text = "\n".join(current_chunk)
            if len(chunk_text) > 50:  # Skip trivially small chunks
                chunks.append({
                    "section": current_section,
                    "text": chunk_text,
                    "start_line": current_start_line,
                })
            
            # Start new chunk
            current_chunk = [line]
            current_start_line = i
            
            # Determine section label
            if match.group(1):
                current_section = match.group(1)
            elif match.group(2):
                current_section = f"{match.group(2)} {match.group(3)}"
            elif match.group(4):
                current_section = f"{match.group(4)} {match.group(5)}"
        else:
            current_chunk.append(line)
    
    # Don't forget the last chunk
    if current_chunk:
        chunk_text = "\n".join(current_chunk)
        if len(chunk_text) > 50:
            chunks.append({
                "section": current_section,
                "text": chunk_text,
                "start_line": current_start_line,
            })
    
    # Merge very small chunks with neighbors (< 100 chars)
    merged = []
    for chunk in chunks:
        if merged and len(chunk["text"]) < 100:
            merged[-1]["text"] += "\n" + chunk["text"]
        else:
            merged.append(chunk)
    
    # Split very large chunks (> 3000 chars) at paragraph boundaries
    final = []
    for chunk in merged:
        if len(chunk["text"]) > 3000:
            # Split at double newlines or sentence boundaries
            sub_texts = re.split(r'\n\s*\n', chunk["text"])
            buffer = ""
            for sub in sub_texts:
                if len(buffer) + len(sub) > 2500 and buffer:
                    final.append({
                        "section": chunk["section"],
                        "text": buffer,
                        "start_line": chunk["start_line"],
                    })
                    buffer = sub
                else:
                    buffer = buffer + "\n" + sub if buffer else sub
            if buffer:
                final.append({
                    "section": chunk["section"],
                    "text": buffer,
                    "start_line": chunk["start_line"],
                })
        else:
            final.append(chunk)
    
    return final


def ingest(file_path: str, name: str = None):
    """Ingest a document into the vector store."""
    import chromadb
    
    file_path = Path(file_path)
    if not file_path.exists():
        print(f"ERROR: File not found: {file_path}")
        return
    
    if name is None:
        name = file_path.stem
    
    # Extract text based on file type
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        data = extract_pdf(str(file_path))
    elif suffix in (".docx", ".doc"):
        data = extract_docx(str(file_path))
    elif suffix in (".txt", ".md"):
        text = file_path.read_text()
        data = {"pages": [{"page_num": 1, "text": text, "tables": []}], "total_pages": 1}
    else:
        print(f"ERROR: Unsupported file type: {suffix}")
        return
    
    # Combine all text
    full_text = ""
    for page in data["pages"]:
        full_text += page["text"] + "\n"
        for table in page["tables"]:
            full_text += "\n[TABLE]\n" + table + "\n[/TABLE]\n"
    
    # Chunk by clauses
    chunks = chunk_by_clauses(full_text, name)
    
    print(f"Document: {name}")
    print(f"Pages: {data['total_pages']}")
    print(f"Total chars: {len(full_text)}")
    print(f"Chunks created: {len(chunks)}")
    
    # Store in chromadb
    os.makedirs(STORE_PATH, exist_ok=True)
    client = chromadb.PersistentClient(path=STORE_PATH)
    
    # Create/get collection
    collection_name = re.sub(r'[^a-zA-Z0-9_]', '_', name.lower())
    
    # Delete existing collection with same name if it exists
    try:
        client.delete_collection(collection_name)
    except Exception:
        pass
    
    collection = client.create_collection(
        name=collection_name,
        metadata={"doc_name": name, "source_file": str(file_path)},
    )
    
    # Add chunks
    ids = []
    documents = []
    metadatas = []
    
    for i, chunk in enumerate(chunks):
        chunk_id = f"{collection_name}_{i:04d}"
        ids.append(chunk_id)
        documents.append(chunk["text"])
        metadatas.append({
            "section": chunk["section"],
            "chunk_index": i,
            "doc_name": name,
            "char_length": len(chunk["text"]),
        })
    
    collection.add(ids=ids, documents=documents, metadatas=metadatas)
    
    # Save index metadata
    index_path = os.path.join(STORE_PATH, "index.json")
    index = {}
    if os.path.exists(index_path):
        with open(index_path) as f:
            index = json.load(f)
    
    index[name] = {
        "collection": collection_name,
        "source_file": str(file_path),
        "chunks": len(chunks),
        "total_chars": len(full_text),
        "pages": data["total_pages"],
        "sections": list(set(c["section"] for c in chunks)),
    }
    
    with open(index_path, "w") as f:
        json.dump(index, f, indent=2)
    
    print(f"Stored in collection: {collection_name}")
    print(f"Sections detected: {len(index[name]['sections'])}")
    print("Done.")


def query(query_text: str, doc_name: str = None, top_k: int = 5):
    """Semantic search across ingested documents."""
    import chromadb
    
    os.makedirs(STORE_PATH, exist_ok=True)
    client = chromadb.PersistentClient(path=STORE_PATH)
    
    # Load index
    index_path = os.path.join(STORE_PATH, "index.json")
    if not os.path.exists(index_path):
        print("ERROR: No documents ingested yet.")
        return
    
    with open(index_path) as f:
        index = json.load(f)
    
    # Determine which collections to search
    if doc_name:
        if doc_name not in index:
            print(f"ERROR: Document '{doc_name}' not found. Available: {list(index.keys())}")
            return
        collections = [index[doc_name]["collection"]]
    else:
        collections = [v["collection"] for v in index.values()]
    
    results = []
    for coll_name in collections:
        try:
            collection = client.get_collection(coll_name)
            res = collection.query(query_texts=[query_text], n_results=top_k)
            
            for i in range(len(res["ids"][0])):
                results.append({
                    "doc": res["metadatas"][0][i].get("doc_name", coll_name),
                    "section": res["metadatas"][0][i].get("section", "?"),
                    "text": res["documents"][0][i],
                    "distance": res["distances"][0][i] if res.get("distances") else None,
                })
        except Exception as e:
            print(f"Warning: Could not query collection {coll_name}: {e}")
    
    # Sort by distance (lower = more relevant)
    results.sort(key=lambda x: x.get("distance") or 999)
    
    # Print results
    print(f"Query: {query_text}")
    print(f"Results: {len(results)}\n")
    
    for i, r in enumerate(results[:top_k], 1):
        print(f"--- Result {i} (doc: {r['doc']}, section: {r['section']}, dist: {r['distance']:.4f}) ---")
        print(r["text"][:500])
        print()


def section_lookup(section_num: str, doc_name: str = None):
    """Look up a specific section by number."""
    import chromadb
    
    os.makedirs(STORE_PATH, exist_ok=True)
    client = chromadb.PersistentClient(path=STORE_PATH)
    
    index_path = os.path.join(STORE_PATH, "index.json")
    if not os.path.exists(index_path):
        print("ERROR: No documents ingested yet.")
        return
    
    with open(index_path) as f:
        index = json.load(f)
    
    if doc_name:
        if doc_name not in index:
            print(f"ERROR: Document '{doc_name}' not found.")
            return
        collections = [(doc_name, index[doc_name]["collection"])]
    else:
        collections = [(k, v["collection"]) for k, v in index.items()]
    
    found = False
    for name, coll_name in collections:
        try:
            collection = client.get_collection(coll_name)
            # Use where clause to filter by section
            res = collection.get(where={"section": section_num})
            
            if res["ids"]:
                found = True
                print(f"=== Section {section_num} (doc: {name}) ===")
                for doc in res["documents"]:
                    print(doc)
                    print()
        except Exception:
            pass
    
    if not found:
        # Try partial match
        print(f"Exact section '{section_num}' not found. Trying semantic search...")
        query(f"section {section_num}", doc_name, top_k=3)


def list_docs():
    """List all ingested documents."""
    index_path = os.path.join(STORE_PATH, "index.json")
    if not os.path.exists(index_path):
        print("No documents ingested yet.")
        return
    
    with open(index_path) as f:
        index = json.load(f)
    
    print(f"Ingested documents ({len(index)}):")
    for name, info in index.items():
        print(f"  - {name}: {info['pages']} pages, {info['chunks']} chunks, {info['total_chars']} chars")


def delete_doc(name: str):
    """Delete a document from the store."""
    import chromadb
    
    index_path = os.path.join(STORE_PATH, "index.json")
    if not os.path.exists(index_path):
        print("No documents ingested.")
        return
    
    with open(index_path) as f:
        index = json.load(f)
    
    if name not in index:
        print(f"Document '{name}' not found. Available: {list(index.keys())}")
        return
    
    client = chromadb.PersistentClient(path=STORE_PATH)
    try:
        client.delete_collection(index[name]["collection"])
    except Exception:
        pass
    
    del index[name]
    with open(index_path, "w") as f:
        json.dump(index, f, indent=2)
    
    print(f"Deleted: {name}")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)
    
    cmd = sys.argv[1]
    
    if cmd == "ingest":
        if len(sys.argv) < 3:
            print("Usage: doc_rag.py ingest <file_path> [--name <name>]")
            sys.exit(1)
        file_path = sys.argv[2]
        name = None
        if "--name" in sys.argv:
            name_idx = sys.argv.index("--name")
            if name_idx + 1 < len(sys.argv):
                name = sys.argv[name_idx + 1]
        ingest(file_path, name)
    
    elif cmd == "query":
        if len(sys.argv) < 3:
            print("Usage: doc_rag.py query <text> [--doc <name>] [--top_k N]")
            sys.exit(1)
        query_text = sys.argv[2]
        doc_name = None
        top_k = 5
        if "--doc" in sys.argv:
            idx = sys.argv.index("--doc")
            doc_name = sys.argv[idx + 1]
        if "--top_k" in sys.argv:
            idx = sys.argv.index("--top_k")
            top_k = int(sys.argv[idx + 1])
        query(query_text, doc_name, top_k)
    
    elif cmd == "section":
        if len(sys.argv) < 3:
            print("Usage: doc_rag.py section <number> [--doc <name>]")
            sys.exit(1)
        section_num = sys.argv[2]
        doc_name = None
        if "--doc" in sys.argv:
            idx = sys.argv.index("--doc")
            doc_name = sys.argv[idx + 1]
        section_lookup(section_num, doc_name)
    
    elif cmd == "list":
        list_docs()
    
    elif cmd == "delete":
        if len(sys.argv) < 3:
            print("Usage: doc_rag.py delete <name>")
            sys.exit(1)
        delete_doc(sys.argv[2])
    
    else:
        print(f"Unknown command: {cmd}")
        print(__doc__)
        sys.exit(1)
